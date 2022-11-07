import logging
import pathlib

from docx import Document
from docx.document import Document as TDocument
from docx.oxml.xmlchemy import _OxmlElementBase
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph, Run
from docx.image.image import Image
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.oxml.ns import qn

from docx_handler.utils import get_paras_by_style_name


def add_logo(doc: TDocument):
    paras = get_paras_by_style_name(doc, "Title")
    if len(paras) != 1:
        logging.error("Title of docx not found")
    title: Paragraph = paras[0]
    logo_path: pathlib.Path = (
        pathlib.Path(__file__).parent.resolve() / "assets" / "hust" / "logo.png"
    )
    image_para: Paragraph = doc.add_paragraph()
    image_para.style = doc.styles["Figure"]
    r = image_para.add_run()
    for _ in range(3):
        r.add_break(WD_BREAK.LINE)
    r.add_picture(str(logo_path), width=Inches(4))
    for _ in range(2):
        r.add_break(WD_BREAK.LINE)
    title._p.addprevious(image_para._p)


def add_student_info_table(doc: TDocument):
    """
    Add the table of student info.

    There will be two columns in the table, the first
    col is title and the second is content.
    
    The first col has no border and the second col will have
    only the bottom border.

    The style of table is `StudentInfoTable`, and the style
    of text is `StudentInfo`
    """
    paras = get_paras_by_style_name(doc, "Subtitle")
    if len(paras) != 1:
        logging.error("Subtitle of docx not found")
    title: Paragraph = paras[0]

    front_break: Paragraph = doc.add_paragraph()
    r: Run = front_break.add_run()
    r.add_break()
    title._p.addprevious(front_break._p)

    back_break: Paragraph = doc.add_paragraph()
    r: Run = back_break.add_run()
    for _ in range(2):
        r.add_break()

    table: Table = doc.add_table(rows=5, cols=2)
    table.style = doc.styles["StudentInfoTable"]
    cell: _Cell

    table.columns[0].cells[0].text = "院系"
    table.columns[0].cells[1].text = "专业班级"
    table.columns[0].cells[2].text = "姓名"
    table.columns[0].cells[3].text = "学号"
    table.columns[0].cells[4].text = "指导老师"

    for cell in table.columns[1].cells:
        cell.width = Inches(3)
        cell.paragraphs[0].style = doc.styles["StudentInfo"]

    for cell in table.columns[0].cells:
        cell.width = Inches(2)
        cell.paragraphs[0].style = doc.styles["StudentInfo"]

    # set table border
    bottom = table._element.xpath("./w:tblPr/w:tblLook")[0]
    bottom.set(qn("w:lastColumn"), "1")
    bottom.set(qn("w:firstRow"), "0")

    title._p.addnext(table._tbl)
    title._p.addnext(back_break._p)

def replace_top_caption(doc:TDocument):
    """
    Replace the caption of "Table of content" by Chinese.

    `sdt` is for "structured document tags", which is the root element TOC.
    The caption of TOC is in the first paragraph of `sdtContent`.
    """

    caption = doc.element.xpath("./w:body/w:sdt/w:sdtContent/w:p/w:r/w:t")[0]
    caption.text = "目录"

def process_hust_docx(filename: str):
    doc: TDocument = Document(filename)
    add_logo(doc)
    add_student_info_table(doc)
    replace_top_caption(doc)
    doc.save(filename)
